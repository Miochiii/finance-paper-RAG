# -*- coding: utf-8 -*-
"""md_docx.py —— 综述 Markdown → Word(docx) 转换 + 引用格式转换 + 论文格式排版。

引用格式（导出时由用户选择）：
  - superscript（上标编号）：[1] 保持数字引用，渲染为 Word 上标；
  - author_year（著者-年份）：[1] → （作者，年份）；相邻多引合并为
    （作者A，年份A；作者B，年份B）；引用若已处于中文括号内则不重复加括号。

排版（依据中南财经政法大学硕士学位论文《排版与印刷要求》，可通过 style 覆盖）：
  - 页面：A4；页边距 上/左 30mm、下/右 25mm；页眉小五宋体（页眉距 15mm）；
    页码小五 Times New Roman、页脚居中（页脚距 15mm）；
  - 章标题（##，大纲 1 级）：黑体三号加粗居中，单倍行距，段前 24 磅、段后 18 磅，
    序号与题名间空两个字符，序号用 Times New Roman；
  - 节标题（###，大纲 2 级）：黑体小三加粗居左，单倍行距，段前 24 磅、段后 6 磅；
  - 目标题（####，大纲 3 级）：宋体四号加粗居左，单倍行距，段前 0.5 行、段后 0；
  - 三级目标题（#####+）：黑体小四，单倍行距，首行缩进 2 字符；
  - 正文：宋体小四（西文 Times New Roman 12 磅），两端对齐，1.5 倍行距，
    首行缩进 2 个汉字字符，段前段后 0；
  - 表格内文字：宋体五号。

编号→作者/年份的映射（ref_map: {n: (author, year)}）：
  1. 优先解析正文末尾"## 参考文献"列表（[N] 作者. 标题（年份））；
  2. 兜底用 survey_export 落盘的 exports/<slug>.map.json + refs.json。
  映射不到的编号保持原样 [N]。
"""

import os
import re
from typing import Dict, List, Optional, Tuple

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

_CITE_RE = re.compile(r"\[(\d+)\]")
_REF_LINE_RE = re.compile(r"^\[(\d+)\]\s*([^\.\s，,]+).*?（(\d{4})）")
_INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
_HEAD_NUM_RE = re.compile(r"^(\d+(?:[\.、]\d+)*)\s+(.+)$")

Segment = Tuple[str, str]  # (kind, text)：text / cite_sup / cite_ay

# --------------------------------------------------------------------------
# 排版默认值（中南财经政法大学硕士学位论文格式）
# --------------------------------------------------------------------------
DEFAULT_STYLE = {
    "page": {
        "top_mm": 30, "left_mm": 30, "bottom_mm": 25, "right_mm": 25,
        "header_mm": 15, "footer_mm": 15,
        "header_text": "中南财经政法大学硕士学位论文",
        "page_number": True,
    },
    "body": {
        "cn_font": "宋体", "en_font": "Times New Roman", "size_pt": 12,
        "line_spacing": 1.5, "first_line_chars": 2,
        "align": "justify", "before_pt": 0, "after_pt": 0,
    },
    "h1": {"cn_font": "黑体", "en_font": "Times New Roman", "size_pt": 16,
           "bold": True, "align": "center", "line_spacing": 1.0,
           "before_pt": 24, "after_pt": 18},
    "h2": {"cn_font": "黑体", "en_font": "Times New Roman", "size_pt": 15,
           "bold": True, "align": "left", "line_spacing": 1.0,
           "before_pt": 24, "after_pt": 6},
    "h3": {"cn_font": "宋体", "en_font": "Times New Roman", "size_pt": 14,
           "bold": True, "align": "left", "line_spacing": 1.0,
           "before_lines": 0.5, "after_pt": 0},
    "h4": {"cn_font": "黑体", "en_font": "Times New Roman", "size_pt": 12,
           "bold": False, "align": "left", "line_spacing": 1.0,
           "first_line_chars": 2, "before_pt": 0, "after_pt": 0},
    "table": {"cn_font": "宋体", "en_font": "Times New Roman", "size_pt": 10.5},
}


def _merge_style(user: Optional[Dict]) -> Dict:
    """用户选项按分组浅合并进默认值（未提供的项保持默认）。"""
    out = {k: dict(v) for k, v in DEFAULT_STYLE.items()}
    for key, val in (user or {}).items():
        if isinstance(val, dict) and key in out:
            out[key].update(val)
        else:
            out[key] = val
    return out


# --------------------------------------------------------------------------
# 引用映射与转换
# --------------------------------------------------------------------------
def extract_ref_map(text: str) -> Dict[int, Tuple[str, str]]:
    """从"## 参考文献"列表解析 [N] → (作者, 年份)。无该节返回空。"""
    m = re.search(r"^##\s*参考文献\s*$", text, re.M)
    if not m:
        return {}
    ref_map: Dict[int, Tuple[str, str]] = {}
    for line in text[m.end():].splitlines():
        s = line.strip()
        if s.startswith("#"):
            break
        if not s:
            continue
        mm = _REF_LINE_RE.match(s)
        if mm:
            ref_map[int(mm.group(1))] = (mm.group(2), mm.group(3))
    return ref_map


def _segments(text: str, ref_map: Dict[int, Tuple[str, str]], fmt: str) -> List[Segment]:
    """把文本切成片段：text（普通文本）/ cite_sup（上标引用）/ cite_ay（著者-年份文本）。"""
    singles = list(_CITE_RE.finditer(text))
    if not singles:
        return [("text", text)]
    segs: List[Segment] = []
    pos, i, n = 0, 0, len(singles)
    while i < n:
        m = singles[i]
        if m.start() > pos:
            segs.append(("text", text[pos:m.start()]))
        # 收集相邻引用（[1][2]）
        group = [m]
        j = i + 1
        while j < n and singles[j].start() == group[-1].end():
            group.append(singles[j])
            j += 1
        i = j
        g_end = group[-1].end()
        # 中文括号内引用（如"（见[1]）"）：引用前有未闭合的（ 且紧随其后是 ）
        # → 著者-年份文本不再加括号，直接替换 [N]
        wrapped = False
        if fmt != "superscript" and text[g_end:g_end + 1] == "）":
            prev = text[pos:m.start()]
            if prev.count("（") > prev.count("）"):
                wrapped = True
        pos = g_end
        if fmt == "superscript":
            segs.append(("cite_sup", text[group[0].start():group[-1].end()]))
        else:  # author_year
            parts = []
            for g in group:
                hit = ref_map.get(int(g.group(1)))
                if hit:
                    parts.append(f"{hit[0]}，{hit[1]}")
            if parts:
                joined = "；".join(parts)
                segs.append(("cite_ay", joined if wrapped else f"（{joined}）"))
            else:
                segs.append(("text", text[group[0].start():group[-1].end()]))
    if pos < len(text):
        segs.append(("text", text[pos:]))
    return segs


def convert_citations(text: str, ref_map: Dict[int, Tuple[str, str]], fmt: str) -> str:
    """纯文本转换（测试/预览用）：返回转换后的整段文本。"""
    return "".join(t for _, t in _segments(text, ref_map, fmt))


# --------------------------------------------------------------------------
# 论文排版（样式/节属性/页眉页码）
# --------------------------------------------------------------------------
def _set_run_font(run, cn_font: str, en_font: str, size_pt: float,
                  bold: Optional[bool] = None):
    run.font.name = en_font
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), en_font)
    rFonts.set(qn("w:hAnsi"), en_font)
    rFonts.set(qn("w:eastAsia"), cn_font)


def _set_spacing(pPr, line: Optional[float] = None, before_pt: Optional[float] = None,
                 after_pt: Optional[float] = None, before_lines: Optional[float] = None):
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        pPr.append(sp)
    if line is not None:
        sp.set(qn("w:line"), str(int(line * 240)))
        sp.set(qn("w:lineRule"), "auto")
    if before_pt is not None:
        sp.set(qn("w:before"), str(int(before_pt * 20)))
    if after_pt is not None:
        sp.set(qn("w:after"), str(int(after_pt * 20)))
    if before_lines is not None:
        sp.set(qn("w:beforeLines"), str(int(before_lines * 100)))


def _set_ind_chars(pPr, chars: int):
    """首行缩进按汉字字符数（w:firstLineChars，100 = 1 字符）。"""
    if chars <= 0:
        return
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(qn("w:firstLineChars"), str(chars * 100))
    ind.set(qn("w:firstLine"), "0")


_ALIGN = {
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _setup_style(doc, style_name: str, spec: Dict):
    """把字体/字号/加粗/对齐/间距/缩进写入 docx 内置样式。"""
    st = doc.styles[style_name]
    st.font.name = spec.get("en_font", "Times New Roman")
    st.font.size = Pt(spec.get("size_pt", 12))
    st.font.bold = spec.get("bold", False)
    st.font.color.rgb = RGBColor(0, 0, 0)  # 去掉模板默认蓝色标题
    rPr = st.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    # 关键：清掉模板的主题字体引用（asciiTheme/eastAsiaTheme 等），
    # 否则 Word 优先用主题字体（默认回退 MS Gothic），显式字体不生效
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        if rFonts.get(qn(attr)) is not None:
            del rFonts.attrib[qn(attr)]
    rFonts.set(qn("w:ascii"), spec.get("en_font", "Times New Roman"))
    rFonts.set(qn("w:hAnsi"), spec.get("en_font", "Times New Roman"))
    rFonts.set(qn("w:eastAsia"), spec.get("cn_font", "宋体"))
    pf = st.paragraph_format
    pf.alignment = _ALIGN.get(spec.get("align", "left"), WD_ALIGN_PARAGRAPH.LEFT)
    pPr = st.element.get_or_add_pPr()
    _set_spacing(pPr, line=spec.get("line_spacing"), before_pt=spec.get("before_pt"),
                 after_pt=spec.get("after_pt"), before_lines=spec.get("before_lines"))
    _set_ind_chars(pPr, spec.get("first_line_chars", 0))


def _setup_section(doc, st: Dict):
    """A4、页边距、页眉/页脚距、页眉文字、页脚页码。"""
    page = st["page"]
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(float(page.get("top_mm", 30)))
    section.left_margin = Mm(float(page.get("left_mm", 30)))
    section.bottom_margin = Mm(float(page.get("bottom_mm", 25)))
    section.right_margin = Mm(float(page.get("right_mm", 25)))
    section.header_distance = Mm(float(page.get("header_mm", 15)))
    section.footer_distance = Mm(float(page.get("footer_mm", 15)))

    header_text = str(page.get("header_text") or "")
    if header_text:
        p = section.header.paragraphs[0]
        p.text = header_text
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            _set_run_font(r, "宋体", "Times New Roman", 9)  # 小五宋体

    if page.get("page_number", True):
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        f1 = OxmlElement("w:fldChar")
        f1.set(qn("w:fldCharType"), "begin")
        it = OxmlElement("w:instrText")
        it.set(qn("xml:space"), "preserve")
        it.text = "PAGE"
        f2 = OxmlElement("w:fldChar")
        f2.set(qn("w:fldCharType"), "end")
        r._r.append(f1)
        r._r.append(it)
        r._r.append(f2)
        _set_run_font(r, "宋体", "Times New Roman", 9)  # 小五 Times New Roman


# --------------------------------------------------------------------------
# Markdown → docx
# --------------------------------------------------------------------------
def _add_text_runs(p, text: str):
    """普通文本段 → docx runs（支持 **粗体** 与 `行内代码`）。"""
    for part in _INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            p.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            p.add_run(part[1:-1]).font.name = "Consolas"
        else:
            p.add_run(part)


def _add_rich_paragraph(p, text: str, ref_map: Dict[int, Tuple[str, str]], fmt: str):
    """段落文本（含引用转换 + 行内格式）→ docx runs。"""
    for kind, seg in _segments(text, ref_map, fmt):
        if kind == "text":
            _add_text_runs(p, seg)
        elif kind == "cite_sup":
            r = p.add_run(seg)
            r.font.superscript = True
        else:  # cite_ay
            p.add_run(seg)


def _add_heading_para(doc, level: int, text: str, spec: Dict):
    """标题段落：序号与题名间空两个字符，序号用西文字体（Times New Roman）。
    run 级显式设置字体（双保险：样式层的主题字体可能覆盖样式属性）。"""
    p = doc.add_paragraph(style=f"Heading {level}")
    t = text.strip()
    m = _HEAD_NUM_RE.match(t)
    if m:
        r1 = p.add_run(m.group(1))
        _set_run_font(r1, spec.get("cn_font", "黑体"), spec.get("en_font", "Times New Roman"),
                      spec.get("size_pt", 16), spec.get("bold", True))
        r2 = p.add_run("　　" + m.group(2))  # 两个全角空格
        _set_run_font(r2, spec.get("cn_font", "黑体"), spec.get("en_font", "Times New Roman"),
                      spec.get("size_pt", 16), spec.get("bold", True))
    else:
        r = p.add_run(t)
        _set_run_font(r, spec.get("cn_font", "黑体"), spec.get("en_font", "Times New Roman"),
                      spec.get("size_pt", 16), spec.get("bold", True))
    return p


def _split_row(line: str) -> List[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def md_to_docx(text: str, out_path: str, title: str = "",
               citation_format: str = "author_year", include_refs: bool = False,
               ref_map: Optional[Dict[int, Tuple[str, str]]] = None,
               style: Optional[Dict] = None) -> None:
    """Markdown → docx。citation_format: superscript / author_year；
    style: 排版覆盖（按 DEFAULT_STYLE 分组浅合并）。"""
    from docx import Document

    st = _merge_style(style)
    doc = Document()

    # 排版：正文/标题样式 + 页面设置
    _setup_style(doc, "Normal", st["body"])
    for level in (1, 2, 3, 4):
        _setup_style(doc, f"Heading {level}", st[f"h{level}"])
    _setup_section(doc, st)

    if title:
        p = doc.add_paragraph(style="Heading 1")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        _set_run_font(r, "黑体", "Times New Roman", 18, True)  # 小二加粗，文档题名

    ref_map = ref_map or {}
    lines = text.replace("\r\n", "\n").split("\n")
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 参考文献节
        if re.match(r"^##\s*参考文献\s*$", stripped):
            if include_refs:
                _add_heading_para(doc, 1, "参考文献", st["h1"])
                i += 1
                while i < n:
                    s = lines[i].strip()
                    if s.startswith("#"):
                        break
                    if s:
                        p = doc.add_paragraph()
                        _add_rich_paragraph(p, s, ref_map, citation_format)
                    i += 1
                continue
            else:
                break  # 不包含参考文献：丢弃其后全部内容

        # 代码块
        if stripped.startswith("```"):
            i += 1
            buf = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            if buf:
                p = doc.add_paragraph()
                r = p.add_run("\n".join(buf))
                r.font.name = "Consolas"
            continue

        # 标题（## → 章，### → 节，#### → 目，更深 → 三级目）
        h = re.match(r"^(#{1,6})\s+(.*)$", line)
        if h:
            level = min(max(1, len(h.group(1)) - 1), 4)
            _add_heading_para(doc, level, h.group(2), st[f"h{level}"])
            i += 1
            continue

        if re.match(r"^(\s*[-*_]){3,}\s*$", line):
            i += 1
            continue

        # 引用块（按正文格式输出）
        if re.match(r"^>\s?", line):
            buf = []
            while i < n and re.match(r"^>\s?", lines[i]):
                buf.append(re.sub(r"^>\s?", "", lines[i]))
                i += 1
            p = doc.add_paragraph()
            _add_rich_paragraph(p, "　".join(buf), ref_map, citation_format)
            continue

        # 列表
        if re.match(r"^\s*[-*+]\s+", line):
            while i < n and re.match(r"^\s*[-*+]\s+", lines[i]):
                p = doc.add_paragraph(style="List Bullet")
                _add_rich_paragraph(p, re.sub(r"^\s*[-*+]\s+", "", lines[i]), ref_map, citation_format)
                i += 1
            continue
        if re.match(r"^\s*\d+[.)]\s+", line):
            while i < n and re.match(r"^\s*\d+[.)]\s+", lines[i]):
                p = doc.add_paragraph(style="List Number")
                _add_rich_paragraph(p, re.sub(r"^\s*\d+[.)]\s+", "", lines[i]), ref_map, citation_format)
                i += 1
            continue

        # 表格（表中文字宋体五号）
        if re.match(r"^\s*\|", line):
            tbl = []
            while i < n and re.match(r"^\s*\|", lines[i]):
                tbl.append(lines[i])
                i += 1
            rows = [_split_row(x) for x in tbl if not re.match(r"^\|[\s:|-]+\|$", x)]
            rows = [r for r in rows if any(c for c in r)]
            if rows:
                ncol = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=ncol)
                table.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci in range(ncol):
                        cell = table.cell(ri, ci)
                        cell.text = row[ci] if ci < len(row) else ""
                        for cp in cell.paragraphs:
                            for cr in cp.runs:
                                _set_run_font(cr, st["table"]["cn_font"],
                                              st["table"]["en_font"],
                                              st["table"]["size_pt"])
            continue

        # 普通段落（合并连续行）
        if stripped:
            buf = [line]
            i += 1
            while i < n and lines[i].strip() and not re.match(
                    r"^(#{1,6}\s|>|\s*[-*+]\s|\s*\d+[.)]\s|\||```|---|\*{3})", lines[i]):
                buf.append(lines[i])
                i += 1
            p = doc.add_paragraph()
            _add_rich_paragraph(p, " ".join(buf), ref_map, citation_format)
            continue
        i += 1

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
