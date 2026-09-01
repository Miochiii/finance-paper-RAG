# -*- coding: utf-8 -*-
"""md_docx 转换器与 Word 导出（引用格式转换）测试。"""
import os

import rag_core.md_docx as M
import rag_core.survey as sv
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def _use_tmp(monkeypatch, work_tmp):
    monkeypatch.setattr(sv, "SURVEY_DIR", os.path.join(work_tmp, "surveys"))


REF_MAP = {1: ("李安哲", "2022"), 2: ("张璐瑶", "2023")}


class TestConvertCitations:
    def test_single(self):
        out = M.convert_citations("……的结果[1]。", REF_MAP, "author_year")
        assert out == "……的结果（李安哲，2022）。"

    def test_adjacent_group(self):
        out = M.convert_citations("多篇研究[1][2]证实。", REF_MAP, "author_year")
        assert out == "多篇研究（李安哲，2022；张璐瑶，2023）证实。"

    def test_unmapped_kept(self):
        assert M.convert_citations("见[9]。", REF_MAP, "author_year") == "见[9]。"

    def test_paren_collapse(self):
        # 引用已在中文括号内 → 不重复加括号
        assert M.convert_citations("（见[1]）", REF_MAP, "author_year") == "（见李安哲，2022）"

    def test_superscript_segments(self):
        segs = M._segments("结论[1]成立", REF_MAP, "superscript")
        assert ("cite_sup", "[1]") in segs
        # 上标模式不产生著者-年份文本
        assert not any(k == "cite_ay" for k, _ in segs)


class TestExtractRefMap:
    def test_parse(self):
        text = ("正文\n\n## 参考文献\n\n"
                "[1] 李安哲. 基于动量反转与机器学习的统计套利策略研究（2022），第 14 页\n"
                "[2] 张璐瑶. 信息熵在机器学习算法中的运用（2023），第 5 页")
        assert M.extract_ref_map(text) == {1: ("李安哲", "2022"), 2: ("张璐瑶", "2023")}

    def test_no_section(self):
        assert M.extract_ref_map("没有参考文献节") == {}


def _write_and_read(tmp, text, **kw):
    p = os.path.join(tmp, "t.docx")
    M.md_to_docx(text, p, title="测试综述", **kw)
    return Document(p)


class TestMdToDocx:
    def test_heading_paragraph_bold_author_year(self, work_tmp):
        doc = _write_and_read(work_tmp, "## 1 引言\n\n这是**加粗**的内容[1]。",
                              ref_map=REF_MAP, citation_format="author_year")
        texts = [p.text for p in doc.paragraphs]
        assert any("引言" in t for t in texts)
        assert any("（李安哲，2022）" in t for t in texts)
        assert any(r.bold for p in doc.paragraphs for r in p.runs)

    def test_superscript(self, work_tmp):
        doc = _write_and_read(work_tmp, "结论[1]成立。",
                              ref_map=REF_MAP, citation_format="superscript")
        assert any(r.font.superscript for p in doc.paragraphs for r in p.runs)

    def test_refs_excluded_by_default(self, work_tmp):
        text = "正文[1]。\n\n## 参考文献\n\n[1] 李安哲. 标题（2022）"
        doc = _write_and_read(work_tmp, text, ref_map=REF_MAP, citation_format="author_year")
        joined = "\n".join(p.text for p in doc.paragraphs)
        assert "参考文献" not in joined

    def test_refs_included(self, work_tmp):
        text = "正文[1]。\n\n## 参考文献\n\n[1] 李安哲. 标题（2022）"
        doc = _write_and_read(work_tmp, text, ref_map=REF_MAP,
                              citation_format="author_year", include_refs=True)
        joined = "\n".join(p.text for p in doc.paragraphs)
        assert "参考文献" in joined and "李安哲. 标题（2022）" in joined


class TestThesisFormatting:
    """论文排版（中南财经政法大学格式默认值）断言。"""

    def test_body_style(self, work_tmp):
        doc = _write_and_read(work_tmp, "正文段落。", ref_map=REF_MAP)
        st = doc.styles["Normal"]
        rfonts = st.element.rPr.rFonts
        assert rfonts.get(qn("w:eastAsia")) == "宋体"
        assert rfonts.get(qn("w:ascii")) == "Times New Roman"
        assert st.font.size.pt == 12
        # 首行缩进 2 字符 + 1.5 倍行距
        assert st.element.pPr.ind.get(qn("w:firstLineChars")) == "200"
        assert st.element.pPr.spacing.get(qn("w:line")) == "360"

    def test_heading_styles(self, work_tmp):
        doc = _write_and_read(work_tmp, "正文。", ref_map=REF_MAP)
        h1 = doc.styles["Heading 1"]
        assert h1.element.rPr.rFonts.get(qn("w:eastAsia")) == "黑体"
        assert h1.font.size.pt == 16 and h1.font.bold is True
        assert h1.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert h1.element.pPr.spacing.get(qn("w:before")) == "480"   # 段前 24 磅
        assert h1.element.pPr.spacing.get(qn("w:after")) == "360"    # 段后 18 磅
        h2 = doc.styles["Heading 2"]
        assert h2.font.size.pt == 15 and h2.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.LEFT
        h3 = doc.styles["Heading 3"]
        assert h3.font.size.pt == 14 and h3.font.bold is True

    def test_page_setup_header_footer(self, work_tmp):
        doc = _write_and_read(work_tmp, "正文。", ref_map=REF_MAP)
        sec = doc.sections[0]
        assert round(sec.page_width.mm) == 210 and round(sec.page_height.mm) == 297
        assert round(sec.top_margin.mm) == 30 and round(sec.left_margin.mm) == 30
        assert round(sec.bottom_margin.mm) == 25 and round(sec.right_margin.mm) == 25
        assert "中南财经政法大学硕士学位论文" in sec.header.paragraphs[0].text
        assert "PAGE" in sec.footer.paragraphs[0]._p.xml  # 页码域

    def test_heading_number_split_no_gap(self, work_tmp):
        doc = _write_and_read(work_tmp, "## 1 引言\n\n正文。", ref_map=REF_MAP)
        para = next(p for p in doc.paragraphs if "引言" in p.text)
        runs = para.runs
        assert runs[0].text == "1"
        assert runs[0].font.name == "Times New Roman"   # 序号用西文字体
        assert runs[1].text == "引言"                    # 序号与题名直连（无全角空格）
        assert "\u3000" not in para.text
        # run 级中文字体显式生效（回归：主题字体 MS Gothic 覆盖问题）
        rPr2 = runs[1]._element.get_or_add_rPr()
        rfonts2 = rPr2.find(qn("w:rFonts"))
        assert rfonts2 is not None and rfonts2.get(qn("w:eastAsia")) == "黑体"

    def test_heading_styles_have_no_theme_fonts(self, work_tmp):
        """回归：样式 rFonts 不得残留主题字体引用（asciiTheme/eastAsiaTheme），
        否则 Word 优先用主题字体（默认回退 MS Gothic），显式字体不生效。"""
        doc = _write_and_read(work_tmp, "正文。", ref_map=REF_MAP)
        for name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
            rf = doc.styles[name].element.rPr.rFonts
            for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
                assert qn(attr) not in rf.attrib, f"{name} 残留 {attr}"

    def test_custom_style_override(self, work_tmp):
        doc = _write_and_read(work_tmp, "正文。", ref_map=REF_MAP,
                              style={"body": {"cn_font": "仿宋", "size_pt": 14}})
        st = doc.styles["Normal"]
        assert st.element.rPr.rFonts.get(qn("w:eastAsia")) == "仿宋"
        assert st.font.size.pt == 14
        # 未覆盖的组保持默认
        assert doc.styles["Heading 1"].font.size.pt == 16

    def test_table_cell_font(self, work_tmp):
        doc = _write_and_read(work_tmp, "| 指标 | 值 |\n|---|---|\n| A | 1 |", ref_map=REF_MAP)
        cell = doc.tables[0].cell(0, 0)
        r = cell.paragraphs[0].runs[0]
        assert r.font.size.pt == 10.5  # 五号


class TestSurveyExportDocx:
    def test_export_author_year_and_sidecar(self, monkeypatch, work_tmp):
        _use_tmp(monkeypatch, work_tmp)
        monkeypatch.setattr(sv, "_collect_evidence", lambda q, top_k=15: [
            {"source": "a.pdf", "page": 3, "author": "甲", "year": 2022, "text": "证据A"}])
        monkeypatch.setattr(sv, "_llm_text", lambda s, u: "正文[来源1]。")
        sv.survey_outline("主题壬", outline=[{"title": "1 引言", "keywords": []}])
        sv.survey_draft("主题壬")
        r = sv.survey_export_docx("主题壬", citation_format="author_year")
        assert r["ok"] and r["filename"].endswith("著者年份.docx")
        assert os.path.isfile(r["path"])
        doc = Document(r["path"])
        joined = "\n".join(p.text for p in doc.paragraphs)
        assert "（甲，2022）" in joined
        # survey_export 已落盘编号顺序侧车
        assert os.path.isfile(os.path.join(sv._survey_dir("主题壬"), "exports", "主题壬.map.json"))

    def test_invalid_format(self, monkeypatch, work_tmp):
        _use_tmp(monkeypatch, work_tmp)
        monkeypatch.setattr(sv, "_collect_evidence", lambda q, top_k=15: [])
        monkeypatch.setattr(sv, "_llm_text", lambda s, u: "正文[来源1]。")
        sv.survey_outline("主题癸", outline=[{"title": "1 引言", "keywords": []}])
        sv.survey_draft("主题癸")
        r = sv.survey_export_docx("主题癸", citation_format="bogus")
        assert r["ok"] is False


class TestDownloadPath:
    def test_whitelist(self, monkeypatch, work_tmp):
        _use_tmp(monkeypatch, work_tmp)
        d = os.path.join(work_tmp, "surveys", "主题子", "exports")
        os.makedirs(d, exist_ok=True)
        ok = os.path.join(d, "a.docx")
        with open(ok, "w") as f:
            f.write("x")
        assert sv.survey_download_path("主题子", "a.docx") == ok
        assert sv.survey_download_path("主题子", "../a.docx") is None
        assert sv.survey_download_path("主题子", "a.exe") is None
        assert sv.survey_download_path("主题子", "不存在.md") is None
